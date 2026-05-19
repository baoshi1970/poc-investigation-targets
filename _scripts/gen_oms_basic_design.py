# -*- coding: utf-8 -*-
"""OMS 基本設計書 (Word) を生成する.

故意の不一致 (実コード vs 設計書):
- DB 接続は Commons DBCP 利用と記載 -> 実装は DBUtil 静的フィールド使い回し
- セッションタイムアウト 30 分と記載 -> web.xml 240 分
- パスワード保存方式: SHA-256 + Salt と記載 -> 実装 MD5 ノーソルト
- 認証失敗 3 回でロックと記載 -> 実装無し
- 文字コードは内部 UTF-8 と記載 -> 実装 Windows-31J
- 例外時はユーザに技術詳細を出さない方針と記載 -> LoginServlet はスタックトレース直出力
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(r"C:\sparticle\poc-investigation-targets\legacy-order-inventory-system\docs\基本設計書.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

# 既定フォント
style = doc.styles["Normal"]
style.font.name = "Yu Gothic"
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Yu Gothic"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    return h


def para(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Yu Gothic"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Yu Gothic"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")


def add_table(headers, rows, widths_cm=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.name = "Yu Gothic"
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
                run.font.bold = True
    for r, row in enumerate(rows, start=1):
        for i, v in enumerate(row):
            t.rows[r].cells[i].text = str(v)
            for p in t.rows[r].cells[i].paragraphs:
                for run in p.runs:
                    run.font.name = "Yu Gothic"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
                    run.font.size = Pt(9.5)
    if widths_cm:
        for r in t.rows:
            for i, w in enumerate(widths_cm):
                r.cells[i].width = Cm(w)


# ===== 表紙 =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("\n\n\n受発注・在庫管理システム (OMS)\n基本設計書")
run.font.name = "Yu Gothic"
run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
run.font.size = Pt(22)
run.bold = True

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("\n\nVer. 3.4  /  2014-11-30")
run.font.name = "Yu Gothic"
run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
run.font.size = Pt(14)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("\n\n\n株式会社サンプル商事\nシステム部")
run.font.name = "Yu Gothic"
run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
run.font.size = Pt(12)

doc.add_page_break()

# ===== 改訂履歴 =====
heading("改訂履歴", 1)
add_table(
    ["版", "改訂日", "改訂者", "改訂内容"],
    [
        ["1.0", "2008-04-15", "田中", "初版"],
        ["2.0", "2011-08-02", "佐藤", "在庫サブシステム追加"],
        ["2.5", "2013-05-30", "佐藤", "受注確定時の在庫引当を楽観排他方式に変更 (T_INVENTORY.VERSION 追加)"],
        ["3.0", "2014-06-15", "田中", "セッション管理方針見直し (タイムアウト 30 分へ短縮)"],
        ["3.4", "2014-11-30", "田中", "パスワード保存方式 SHA-256 + Salt へ移行"],
    ],
    widths_cm=[1.5, 2.5, 2.5, 10],
)

doc.add_page_break()

# ===== 1. はじめに =====
heading("1. はじめに", 1)
heading("1.1 本書の位置付け", 2)
para("本書は受発注・在庫管理システム (以下「OMS」) の基本設計を定義するものである。")
para("基本設計は、要件定義書 (REQ-2008-OMS-v1.3) の内容を受け、システム化方針および主要機能の構成を明示する。")

heading("1.2 用語定義", 2)
add_table(
    ["用語", "意味"],
    [
        ["受注", "取引先から商品の購入依頼を受け、社内に登録した時点の業務単位"],
        ["引当",  "受注確定に伴い、対応する在庫数を割り当てる処理"],
        ["論理削除", "DEL_FLG='1' を立てるのみで物理削除しない方式"],
    ],
    widths_cm=[3, 13],
)

# ===== 2. システム概要 =====
heading("2. システム概要", 1)
heading("2.1 システム構成", 2)
para("Web 3 層構成。プレゼンテーション層に Tomcat、データ層に Oracle 10g を採用する。")
para("クライアントは社内 LAN 内の Windows PC のみを想定し、ブラウザは Internet Explorer 8 以上を必須とする。")

heading("2.2 主要構成要素", 2)
add_table(
    ["階層", "構成要素", "採用技術"],
    [
        ["プレゼンテーション層", "JSP / Servlet (一部 Struts1 タグ)", "JSP 2.1, Servlet 2.5"],
        ["業務ロジック層", "POJO Servlet + 自前 DAO", "Java 1.6"],
        ["データアクセス層", "Commons DBCP によるコネクションプール経由 JDBC", "Commons DBCP 1.2"],
        ["DB",            "Oracle 10g Standard Edition", "10.2.0.4"],
        ["Web サーバ",    "Apache Tomcat 6.0",            "6.0.x"],
    ],
    widths_cm=[4, 6, 6],
)

# ===== 3. 方式設計 =====
heading("3. 方式設計", 1)
heading("3.1 セッション管理", 2)
para("認証成功時にセッションを払い出し、HttpSession に USER_ID / USER_NM / DEPT_CD を格納する。")
para("セッションタイムアウトは 30 分とし、無操作で 30 分を超えた場合は再ログインを要求する。")
para("セッション固定攻撃対策として、認証成功直後に session.invalidate() → 再採番を実施する。",
     italic=True, color=RGBColor(0x80, 0x80, 0x80))

heading("3.2 認証", 2)
para("パスワードは SHA-256 + ユーザ毎の固定 Salt でハッシュ化して保管する。")
para("認証失敗が 3 回連続した場合、M_USER.FAIL_CNT をインクリメントし、3 回到達で LOCK_FLG='1' とする。")
para("ロック解除は管理者の手動操作、または夜間バッチ BatchUnlockUsers で 02:00 に日次クリアする。")

heading("3.3 DB アクセス", 2)
para("DB 接続は Commons DBCP 1.2 によるコネクションプールを使用する。")
para("プール定義:")
add_table(
    ["パラメータ", "値"],
    [
        ["initialSize",            "5"],
        ["maxActive",              "30"],
        ["maxIdle",                "10"],
        ["validationQuery",        "SELECT 1 FROM DUAL"],
        ["testOnBorrow",           "true"],
    ],
    widths_cm=[5, 10],
)
para("業務処理は 1 リクエスト = 1 トランザクションを原則とし、Servlet 入口で取得・出口でコミット／ロールバックする。",
     italic=True)

heading("3.4 在庫引当の排他制御", 2)
para("受注確定時に T_INVENTORY.STOCK_QTY を更新する処理は、楽観排他で実装する。")
para("UPDATE 文には WHERE 句に VERSION 列の一致条件を必須化し、更新件数が 0 件の場合は再読込→再試行とする (最大 3 回)。")

heading("3.5 文字コード", 2)
para("内部処理および DB は UTF-8 を採用する。")
para("外部 I/F (取引先からの CSV 取込) のみ Shift_JIS で受け入れ、入口で UTF-8 に変換する。")

heading("3.6 例外処理", 2)
para("業務例外はログに記録の上、ユーザにはエラーコードとガイダンスのみ表示する。スタックトレース等の技術情報は画面に出力しない。")
para("システム例外は web.xml の <error-page> 経由で SCR-099 エラー画面に遷移させる。")

# ===== 4. バッチ設計 =====
heading("4. バッチ設計", 1)
add_table(
    ["バッチID", "バッチ名", "起動時刻", "概要"],
    [
        ["BAT-001", "BatchUnlockUsers",      "02:00", "M_USER の FAIL_CNT, LOCK_FLG を日次クリア"],
        ["BAT-002", "BatchRebuildCache",     "03:00", "T_INVENTORY_CACHE を再構築"],
        ["BAT-010", "BatchOrderArchive",     "04:00", "確定後 1 年以上経過した受注をアーカイブ DB へ移動"],
    ],
    widths_cm=[2, 4, 2, 8],
)

# ===== 5. 非機能要件 =====
heading("5. 非機能要件", 1)
add_table(
    ["分類", "要件", "目標値"],
    [
        ["性能",        "受注一覧表示の応答時間 (1万件以下)",   "3 秒以内"],
        ["性能",        "在庫照会の応答時間",                    "2 秒以内"],
        ["可用性",      "営業時間内の稼働率",                    "99.5% 以上"],
        ["保守",        "Tomcat 再起動を要する障害発生頻度",      "月 1 回以下"],
        ["セキュリティ","パスワード保存方式",                    "SHA-256 + ユーザ毎 Salt"],
        ["セキュリティ","SQL 実行方式",                          "PreparedStatement 必須"],
        ["セキュリティ","セッションタイムアウト",                "30 分"],
    ],
    widths_cm=[3, 9, 5],
)

# ===== 6. 制約事項 =====
heading("6. 制約事項", 1)
bullet("対応ブラウザは IE8 以上のみ（Chrome / Firefox は対象外）")
bullet("クライアント PC は社内 LAN 内に限る")
bullet("対応 OS は Windows XP SP3 / Windows 7 のみ")
bullet("音声・動画コンテンツは対応しない")

# ===== 付録 =====
doc.add_page_break()
heading("付録 A. 想定外事項 (申し送り)", 1)
para("以下、初期設計時の懸念事項。後続案件にて対応を検討すること。", italic=True)
bullet("受注番号の払い出しは 6 桁連番（日次クリア）のため、1 日 100 万件超でオーバーフローする。現状の業務量なら問題なし。")
bullet("取引先名の半角カナ・機種依存文字は運用ルールで禁止としているが、入力チェックは未実装。")
bullet("在庫 CSV 取込は将来要件。現バージョンでは未対応。")

doc.save(OUT)
print("OK")
