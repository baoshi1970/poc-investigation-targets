# -*- coding: utf-8 -*-
"""Asteria 基本設計書 (Word).

故意の不一致:
- ETL 方針「差分更新を基本とする」と記載 -> 実装は全件 TRUNCATE → BULK INSERT
- リトライ 3 回 + バックオフと記載 -> 実装に該当無し
- 接続情報は Warp の暗号化機能で保管と記載 -> 実は jdbc_connections.xml に平文
- 監視に Zabbix-Agent 連携と記載 -> 実装は SMTP のみ
- LookupDB のキャッシュは ON 必須と記載 -> 実装は cache=false (TICKET-ASTR-012 の原因)
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT = Path(r"C:\sparticle\poc-investigation-targets\asteria-etl-sample\docs\基本設計書.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Yu Gothic"
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Yu Gothic"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    return h


def para(text, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Yu Gothic"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    run.italic = italic
    if color:
        run.font.color.rgb = color


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Yu Gothic"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")


def add_table(headers, rows, widths_cm=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.font.name = "Yu Gothic"
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
                run.font.bold = True
    for r, row in enumerate(rows, start=1):
        for i, v in enumerate(row):
            t.rows[r].cells[i].text = str(v)
            for p in t.rows[r].cells[i].paragraphs:
                for run in p.runs:
                    run.font.name = "Yu Gothic"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
                    run.font.size = Pt(9.5)
    if widths_cm:
        for r in t.rows:
            for i, w in enumerate(widths_cm):
                r.cells[i].width = Cm(w)


# 表紙
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("\n\n\nOMS-DWH 連携 ETL\n基本設計書")
run.font.name = "Yu Gothic"
run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
run.font.size = Pt(22)
run.bold = True

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("\n\nVer. 1.1  /  2019-10-01")
run.font.name = "Yu Gothic"
run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
run.font.size = Pt(14)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("\n\n\n株式会社サンプル商事\nシステム部 / 外部 SI ベンダー (㈱○○ソリューションズ)")
run.font.name = "Yu Gothic"
run.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
run.font.size = Pt(11)

doc.add_page_break()

heading("改訂履歴", 1)
add_table(
    ["版", "改訂日", "改訂者", "改訂内容"],
    [
        ["1.0", "2017-06-15", "外部SI", "初版 (ASTERIA Warp 1709 導入)"],
        ["1.1", "2019-10-01", "鈴木",   "軽減税率対応 (TAX_RATE 列追加 + XSLT 改修)"],
    ],
    widths_cm=[1.5, 2.5, 2.5, 10],
)

doc.add_page_break()

heading("1. はじめに", 1)
heading("1.1 本書の目的", 2)
para("本書は ASTERIA Warp を用いた OMS-DWH 連携 ETL の基本設計を定義する。")
para("ETL は社内基幹 (OMS) と DWH の間のデータ連携を担う重要システムであり、停止すると翌朝の経営ダッシュボードに影響する。")

heading("1.2 対象範囲", 2)
bullet("取引先マスタの DWH への連携 (FLW-001)")
bullet("受注ヘッダの DWH への連携 (FLW-002)")
bullet("在庫マスタの突合・差分レポート (FLW-003)")

heading("2. システム構成", 1)
heading("2.1 構成図", 2)
para("[構成図はここに挿入] (Visio 2010 ファイル: \\\\fileserver\\dwh\\設計\\構成図.vsd)")
para("構成: OMS(Oracle) ──Asteria Warp──> ファイル共有 ──> DWH(SQL Server)")

heading("2.2 採用基盤", 2)
add_table(
    ["項目", "内容"],
    [
        ["ETL 製品",  "ASTERIA Warp 1709"],
        ["実行基盤",  "専用 Windows Server 2012 R2 (8 vCPU / 16GB RAM)"],
        ["スケジューラ", "Warp 内蔵スケジューラ"],
        ["接続情報管理", "Warp 設定の暗号化機能で保管"],
        ["ファイル中継", "CIFS 共有 (\\\\fileserver01\\etl_stage)"],
        ["DWH",        "Microsoft SQL Server 2016"],
    ],
    widths_cm=[5, 11],
)

heading("3. ETL 方針", 1)
heading("3.1 更新方式", 2)
para("DWH 側テーブルの更新は、原則として「差分更新方式」を採用する。")
para("受注 (FLW-002) は前日分のみを取得し、ファクトテーブル FACT_ORDER に MERGE で投入する。")
para("マスタ系 (FLW-001) も差分検知の上、変更レコードのみ MERGE で更新する。",
     italic=True, color=RGBColor(0x55, 0x55, 0x55))

heading("3.2 トランザクション境界", 2)
para("各フローは 1 フロー = 1 トランザクションを原則とし、ファイル出力 → DWH 取込のいずれかが失敗した場合はロールバックする。")
para("DWH 側で取り込みが完了しなかった場合、Asteria 側で出力したファイルは破棄して翌実行時にリトライする。")

heading("3.3 LookupDB のキャッシュ", 2)
para("各フローで取引先名等を引き当てる LookupDB コンポーネントは、キャッシュ ON を必須とする (cache=true)。")
para("理由: 同一キーで複数回 SQL を発行することによるパフォーマンス劣化を回避するため。")

heading("3.4 エラーハンドリング", 2)
para("フロー実行中の異常発生時は、Warp スケジューラのリトライ機能で 3 回まで自動リトライを行う (バックオフ: 5 分 / 15 分 / 30 分)。")
para("3 回失敗時は it-ops にメール通知し、フローを停止する。")

heading("3.5 文字コード", 2)
para("Asteria 内部処理および DWH との連携ファイルは UTF-8 (BOM 無し) を採用する。")
para("DWH 側 BULK INSERT には CODEPAGE='65001' を明示する。")

heading("4. 監視", 1)
para("フロー終了時に Zabbix-Agent に対しメトリクスを送信する。")
para("メトリクス: フローID / 開始時刻 / 終了時刻 / 処理件数 / 戻り値")
para("Zabbix 側で「3 回連続失敗」または「想定所要時間の 2 倍を超過」を検知した場合、システム部宛にアラート通知する。")

heading("5. セキュリティ", 1)
bullet("DB 接続情報は Warp の暗号化機能で保管し、平文での XML 直書きは禁止する。")
bullet("ETL 専用アカウント (etl_reader / etl_user) は読み取り権限のみとする。DML 権限は付与しない。")
bullet("ファイル共有上のステージングファイルは取込完了後 24 時間で自動削除する。")
bullet("Warp Designer へのアクセスはシステム部の専用 PC のみとし、踏み台経由でのリモート操作は禁止する。")

heading("6. 性能要件", 1)
add_table(
    ["フロー", "想定所要時間", "件数想定", "備考"],
    [
        ["FLW-001 customer_master_sync", "10 分以内", "8 万件",   "差分なら 1 分以内"],
        ["FLW-002 order_etl_daily",      "60 分以内", "前日分 1 万件", "差分更新のため件数増に強い"],
        ["FLW-003 inventory_reconcile",  "90 分以内", "5 万件",   "倉庫毎並列化を推奨"],
    ],
    widths_cm=[6, 4, 3, 4],
)

heading("7. 運用ルール", 1)
bullet("障害時の一次対応はシステム部 鈴木が担当 (兼任)")
bullet("外部 SI ベンダー (㈱○○ソリューションズ) は契約上、メジャー改修時のみ稼働")
bullet("Warp のバージョンアップは年 1 回 (1 月) に検討")
bullet("ステージングファイルの世代管理は DWH 側で実施 (OMS 側は出力のみ)")

heading("付録 A. 課題・将来対応", 1)
para("運用開始時点で以下の課題が認識されている。将来案件で対応する。", italic=True)
bullet("倉庫追加時のフロー分岐をテンプレート化して手動コピペを避ける (現在は手作業)")
bullet("XSLT 内の税率判定ロジックの一元化 (現在は備考欄文字列ベースで脆弱)")
bullet("DWH 側ファクトテーブルのパーティション戦略の見直し")

doc.save(OUT)
print("OK")
